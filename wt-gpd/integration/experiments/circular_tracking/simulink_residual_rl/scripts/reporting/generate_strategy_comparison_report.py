# -*- coding: utf-8 -*-
"""Generate multi-controller circle disturbance comparison report."""

from __future__ import annotations

import csv
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
REPORT_DIR = ROOT / "reports" / "strategy_comparison"
FIG_DIR = ROOT / "results" / "figures"
DATA_DIR = ROOT / "results" / "data"


def set_run_font(run, size=11, bold=False, color=None, name="Microsoft YaHei"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)


def add_para(doc, text="", size=11, bold=False, color=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=14, bold=True, color=(31, 78, 121))
    return p


def add_picture(doc, image_name, caption, width=6.2):
    image_path = FIG_DIR / image_name
    if not image_path.exists():
        add_para(doc, f"[缺失图片: {image_name}]", color=(180, 0, 0))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, size=9, color=(102, 102, 102))


def read_metrics():
    path = DATA_DIR / "quadrotor_strategy_circle_comparison_metrics.csv"
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def f(row, key):
    return float(row[key])


def lookup(rows, model, controller):
    for row in rows:
        if row["model_type"] == model and row["controller_type"] == controller:
            return row
    raise KeyError((model, controller))


def reduction(base, candidate, key):
    b = f(base, key)
    c = f(candidate, key)
    if abs(b) < 1e-12:
        return 0.0
    return 100.0 * (b - c) / b


def metrics_table_markdown(rows):
    lines = [
        "| 模型 | 控制器 | 全程RMS/m | 稳定段RMS/m | 终端误差/m | 最大高度误差/m | 控制努力 |",
        "|---|---:|---:|---:|---:|---:|---:|",
    ]
    for row in rows:
        lines.append(
            f"| {row['model_label']} | {row['controller_label']} | "
            f"{f(row, 'rms_position_error_m'):.4f} | "
            f"{f(row, 'steady_rms_position_error_m'):.4f} | "
            f"{f(row, 'final_position_error_m'):.4f} | "
            f"{f(row, 'max_altitude_error_m'):.4f} | "
            f"{f(row, 'mean_control_effort_rad_s'):.4f} |"
        )
    return "\n".join(lines)


def add_metrics_table(doc, rows):
    headers = ["模型", "控制器", "全程RMS/m", "稳定段RMS/m", "终端误差/m", "最大高度误差/m", "控制努力"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, title in enumerate(headers):
        run = table.rows[0].cells[i].paragraphs[0].add_run(title)
        set_run_font(run, size=7, bold=True)
    for row in rows:
        vals = [
            row["model_label"],
            row["controller_label"],
            f"{f(row, 'rms_position_error_m'):.4f}",
            f"{f(row, 'steady_rms_position_error_m'):.4f}",
            f"{f(row, 'final_position_error_m'):.4f}",
            f"{f(row, 'max_altitude_error_m'):.4f}",
            f"{f(row, 'mean_control_effort_rad_s'):.4f}",
        ]
        cells = table.add_row().cells
        for i, val in enumerate(vals):
            run = cells[i].paragraphs[0].add_run(val)
            set_run_font(run, size=7)


def write_markdown(rows):
    temp_base = lookup(rows, "temperature", "baseline")
    temp_pidff = lookup(rows, "temperature", "pid_ff")
    temp_mpc = lookup(rows, "temperature", "mpc")
    temp_adrc = lookup(rows, "temperature", "adrc")
    temp_rl = lookup(rows, "temperature", "rl")
    dust_base = lookup(rows, "dust", "baseline")
    dust_pidff = lookup(rows, "dust", "pid_ff")
    dust_mpc = lookup(rows, "dust", "mpc")
    dust_adrc = lookup(rows, "dust", "adrc")
    dust_rl = lookup(rows, "dust", "rl")
    std_mpc = lookup(rows, "standard", "mpc")
    std_adrc = lookup(rows, "standard", "adrc")

    lines = [
        "# 多控制策略圆周抗扰对比报告",
        "",
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}",
        "",
        "## 1. 对比目标",
        "",
        "本报告在同一四旋翼 Simulink 模型和同一匀速圆周参考轨迹下，对比原 PID、PID 扰动补偿、线性 MPC、ADRC 和强化学习残差策略。三类环境为标准模型、温度扰动模型和粉尘扰动模型。",
        "",
        "## 2. 控制策略说明",
        "",
        "- 原 PID：当前项目已有串级 PID/PD 控制器，不使用环境扰动信息。",
        "- PID 扰动补偿：保留原 PID 外环和姿态内环，显式补偿风阻、热上升、低密度和粉尘效率折减。",
        "- 线性 MPC：用 18 步有限时域双积分预测外环生成加速度指令，并加入加速度约束，再叠加同一扰动补偿。",
        "- ADRC：以位置/速度双积分模型为对象，使用线性扩张状态观测器估计总扰动，并对执行器效率下降进行推力/力矩调度。",
        "- 强化学习策略：保留低层稳定控制器，由已训练残差策略输出环境补偿动作。",
        "",
        "## 3. 指标汇总",
        "",
        metrics_table_markdown(rows),
        "",
        "## 4. 主要结论",
        "",
        f"- 温度扰动下，原 PID 全程 RMS 为 {f(temp_base, 'rms_position_error_m'):.4f} m；PID 补偿降至 {f(temp_pidff, 'rms_position_error_m'):.4f} m，下降 {reduction(temp_base, temp_pidff, 'rms_position_error_m'):.1f}%；MPC、ADRC、RL 分别为 {f(temp_mpc, 'rms_position_error_m'):.4f} m、{f(temp_adrc, 'rms_position_error_m'):.4f} m、{f(temp_rl, 'rms_position_error_m'):.4f} m。",
        f"- 粉尘扰动下，原 PID 全程 RMS 为 {f(dust_base, 'rms_position_error_m'):.4f} m；PID 补偿、MPC、ADRC、RL 分别为 {f(dust_pidff, 'rms_position_error_m'):.4f} m、{f(dust_mpc, 'rms_position_error_m'):.4f} m、{f(dust_adrc, 'rms_position_error_m'):.4f} m、{f(dust_rl, 'rms_position_error_m'):.4f} m。",
        f"- MPC 在温度扰动下取得最低全程 RMS 和终端误差，但控制努力高于 PID 补偿和 RL；标准模型下 MPC RMS 为 {f(std_mpc, 'rms_position_error_m'):.4f} m，说明它更积极地追赶圆周初始速度。",
        f"- ADRC 标准模型 RMS 为 {f(std_adrc, 'rms_position_error_m'):.4f} m；它不依赖风场/热上升的精确补偿公式，主要优势是用 ESO 在线估计总扰动。",
        f"- RL 策略不需要手写完整环境补偿公式，温度和粉尘下表现接近 PID 补偿；粉尘最大高度误差为 {f(dust_rl, 'max_altitude_error_m'):.4f} m，是五种方法中的重要优势指标。",
        "",
        "## 5. 图表文件",
        "",
        "- `results/figures/strategy_circle_trajectory_3d.png`：三类环境下四种控制器三维轨迹。",
        "- `results/figures/strategy_circle_position_error.png`：位置误差时序。",
        "- `results/figures/strategy_circle_metric_bars.png`：全程 RMS、稳定段 RMS、终端误差柱状图。",
        "- `results/figures/strategy_circle_effort_altitude.png`：控制努力和最大高度误差对比。",
        "",
        "## 6. 结论建议",
        "",
        "从工程部署角度，PID 扰动补偿最易解释且效果已经明显；MPC 在温度扰动全程 RMS 上通常最强，但控制努力更高；ADRC 提供不依赖精确扰动模型的传统抗扰基线；RL 策略在不牺牲标准工况的前提下获得接近模型补偿的抗扰效果，适合作为复杂未知扰动下的残差补偿路线。",
    ]
    (REPORT_DIR / "多控制策略圆周抗扰对比报告.md").write_text("\n".join(lines), encoding="utf-8-sig")


def write_docx(rows):
    temp_base = lookup(rows, "temperature", "baseline")
    temp_mpc = lookup(rows, "temperature", "mpc")
    temp_adrc = lookup(rows, "temperature", "adrc")
    temp_rl = lookup(rows, "temperature", "rl")
    dust_rl = lookup(rows, "dust", "rl")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("多控制策略圆周抗扰对比报告")
    set_run_font(run, size=18, bold=True, color=(31, 78, 121))
    add_para(doc, f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}", size=9, color=(102, 102, 102), align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "1. 控制策略")
    add_para(doc, "本报告比较原 PID、PID 扰动补偿、线性 MPC、ADRC 和强化学习残差策略。所有方法共用同一四旋翼动力学、同一圆周轨迹和同一温度/粉尘扰动模型。")
    add_para(doc, "PID 补偿和 MPC 都显式使用环境模型输出；ADRC 使用扩张状态观测器估计总扰动，并对执行器效率下降做推力/力矩调度；强化学习策略使用训练得到的残差权重；原 PID 不使用环境量。")

    add_heading(doc, "2. 指标汇总")
    add_metrics_table(doc, rows)
    add_para(doc, f"温度扰动下，原 PID RMS 为 {f(temp_base, 'rms_position_error_m'):.4f} m，MPC 降至 {f(temp_mpc, 'rms_position_error_m'):.4f} m，ADRC 为 {f(temp_adrc, 'rms_position_error_m'):.4f} m，RL 为 {f(temp_rl, 'rms_position_error_m'):.4f} m。", bold=True, color=(31, 78, 121))
    add_para(doc, f"粉尘扰动下，RL 最大高度误差为 {f(dust_rl, 'max_altitude_error_m'):.4f} m，显示残差补偿能有效修正推力效率下降。", bold=True, color=(31, 78, 121))

    add_heading(doc, "3. 图表")
    add_picture(doc, "strategy_circle_trajectory_3d.png", "图1 多控制策略匀速圆周三维轨迹对比")
    add_picture(doc, "strategy_circle_position_error.png", "图2 多控制策略位置误差时序对比")
    add_picture(doc, "strategy_circle_metric_bars.png", "图3 RMS 与终端误差指标对比")
    add_picture(doc, "strategy_circle_effort_altitude.png", "图4 控制努力与最大高度误差对比")

    add_heading(doc, "4. 结论")
    add_para(doc, "PID 扰动补偿是最容易解释和部署的增强基线；MPC 在温度扰动下轨迹指标强，但控制努力更高；ADRC 提供不依赖精确扰动模型的传统抗扰基线；强化学习策略在保持标准工况表现的同时，在温度和粉尘扰动下接近模型补偿效果，适合用于未知扰动残差补偿。")

    doc.save(REPORT_DIR / "多控制策略圆周抗扰对比报告.docx")


def main():
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    rows = read_metrics()
    write_markdown(rows)
    write_docx(rows)
    print(REPORT_DIR / "多控制策略圆周抗扰对比报告.md")
    print(REPORT_DIR / "多控制策略圆周抗扰对比报告.docx")


if __name__ == "__main__":
    main()
