# -*- coding: utf-8 -*-
"""Generate RL-v2 vs MPC benchmark report."""

from __future__ import annotations

import csv
import subprocess
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
REPORT_DIR = ROOT / "reports" / "rl_v2"
FIG_DIR = ROOT / "results" / "figures"
DATA_DIR = ROOT / "results" / "data"
STEM = "RL-v2超越MPC控制策略报告"


def set_run_font(run, size=11, bold=False, color=None, name="Microsoft YaHei"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)


def add_para(doc, text="", size=11, bold=False, color=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=14, bold=True, color=(31, 78, 121))


def add_picture(doc, image_name, caption, width=6.2):
    image_path = FIG_DIR / image_name
    if not image_path.exists():
        add_para(doc, f"[缺失图片: {image_name}]", color=(180, 0, 0))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, size=9, color=(102, 102, 102))


def read_metrics():
    path = DATA_DIR / "quadrotor_rl_v2_mpc_benchmark_metrics.csv"
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def f(row, key):
    return float(row[key])


def lookup(rows, model, controller):
    for row in rows:
        if row["model_type"] == model and row["controller_type"] == controller:
            return row
    raise KeyError((model, controller))


def pass_fail(rows):
    std_mpc = lookup(rows, "standard", "mpc")
    std_rl2 = lookup(rows, "standard", "rl_v2")
    temp_mpc = lookup(rows, "temperature", "mpc")
    temp_rl2 = lookup(rows, "temperature", "rl_v2")
    dust_mpc = lookup(rows, "dust", "mpc")
    dust_rl2 = lookup(rows, "dust", "rl_v2")

    checks = [
        ("温度全程 RMS < MPC", f(temp_rl2, "rms_position_error_m") < f(temp_mpc, "rms_position_error_m")),
        ("温度全程 RMS < 0.0877 m", f(temp_rl2, "rms_position_error_m") < 0.0877),
        (
            "温度稳定段 RMS < MPC 或综合代价 < MPC",
            f(temp_rl2, "steady_rms_position_error_m") < f(temp_mpc, "steady_rms_position_error_m")
            or f(temp_rl2, "composite_cost") < f(temp_mpc, "composite_cost"),
        ),
        (
            "温度终端误差 <= 1.1 x MPC",
            f(temp_rl2, "final_position_error_m") <= 1.1 * f(temp_mpc, "final_position_error_m"),
        ),
        ("标准 RMS <= MPC", f(std_rl2, "rms_position_error_m") <= f(std_mpc, "rms_position_error_m")),
        ("粉尘 RMS <= MPC", f(dust_rl2, "rms_position_error_m") <= f(dust_mpc, "rms_position_error_m")),
        ("最大倾角 < 0.35 rad", max(f(r, "max_tilt_command_rad") for r in rows if r["controller_type"] == "rl_v2") < 0.35),
        (
            "控制努力 <= 1.2 x MPC",
            all(
                f(lookup(rows, model, "rl_v2"), "mean_control_effort_rad_s")
                <= 1.2 * f(lookup(rows, model, "mpc"), "mean_control_effort_rad_s")
                for model in ["standard", "temperature", "dust"]
            ),
        ),
        ("旋翼命令饱和率为 0", max(f(r, "rotor_saturation_rate") for r in rows if r["controller_type"] == "rl_v2") <= 1e-12),
    ]
    return checks


def metrics_table_markdown(rows):
    lines = [
        "| 模型 | 控制器 | RMS/m | 稳定RMS/m | 终端/m | 高度/m | 努力 | 倾角/rad | 饱和率 | 综合代价 |",
        "|---|---:|---:|---:|---:|---:|---:|---:|---:|---:|",
    ]
    for row in rows:
        lines.append(
            f"| {row['model_label']} | {row['controller_label']} | "
            f"{f(row, 'rms_position_error_m'):.4f} | "
            f"{f(row, 'steady_rms_position_error_m'):.4f} | "
            f"{f(row, 'final_position_error_m'):.4f} | "
            f"{f(row, 'max_altitude_error_m'):.4f} | "
            f"{f(row, 'mean_control_effort_rad_s'):.4f} | "
            f"{f(row, 'max_tilt_command_rad'):.4f} | "
            f"{f(row, 'rotor_saturation_rate'):.4f} | "
            f"{f(row, 'composite_cost'):.4f} |"
        )
    return "\n".join(lines)


def write_markdown(rows):
    checks = pass_fail(rows)
    temp_mpc = lookup(rows, "temperature", "mpc")
    temp_rl2 = lookup(rows, "temperature", "rl_v2")
    overall = all(ok for _, ok in checks)

    lines = [
        "# RL-v2超越MPC控制策略报告",
        "",
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}",
        "",
        "## 1. 目标",
        "",
        "本报告比较原PID、PID扰动补偿、线性MPC、ADRC、RL-v1 和 RL-v2 六种控制策略，重点验证 RL-v2 是否在匀速圆周抗扰任务中达到优于 MPC 的控制效果。",
        "",
        "## 2. RL-v2 方法",
        "",
        "RL-v2 保留原稳定底层控制器，在外环加入前瞻残差补偿。策略输入包含当前误差、环境扰动、旋翼余量和未来4个参考点；输出为加速度残差、推力修正和力矩修正。训练流程包含 MPC imitation 和 CEM 直接策略搜索，训练过程会打开实时进度窗口显示 cost、RMS 和 MPC 基准。",
        "",
        "## 3. 指标汇总",
        "",
        metrics_table_markdown(rows),
        "",
        "## 4. 验收结果",
        "",
    ]
    for name, ok in checks:
        lines.append(f"- {'PASS' if ok else 'FAIL'}：{name}")
    lines += [
        "",
        f"总体结论：{'RL-v2 已达到优于 MPC 的验收标准。' if overall else 'RL-v2 尚未完全达到优于 MPC 的验收标准，需要继续训练或调参。'}",
        "",
        f"温度扰动下，MPC RMS 为 {f(temp_mpc, 'rms_position_error_m'):.4f} m，RL-v2 RMS 为 {f(temp_rl2, 'rms_position_error_m'):.4f} m。",
        "",
        "## 5. 图表文件",
        "",
        "- `results/figures/rl_v2_benchmark_trajectory_3d.png`：六控制器三维轨迹对比。",
        "- `results/figures/rl_v2_benchmark_position_error.png`：位置误差时序。",
        "- `results/figures/rl_v2_benchmark_metric_bars.png`：RMS、稳定段 RMS、终端误差。",
        "- `results/figures/rl_v2_benchmark_effort_feasibility.png`：控制努力、倾角、旋翼饱和率。",
    ]
    (REPORT_DIR / f"{STEM}.md").write_text("\n".join(lines), encoding="utf-8-sig")


def add_metrics_table(doc, rows):
    headers = ["模型", "控制器", "RMS/m", "稳定RMS/m", "终端/m", "高度/m", "努力", "倾角", "饱和率"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, title in enumerate(headers):
        run = table.rows[0].cells[i].paragraphs[0].add_run(title)
        set_run_font(run, size=7, bold=True)
    for row in rows:
        vals = [
            row["model_label"],
            row["controller_label"],
            f"{f(row, 'rms_position_error_m'):.4f}",
            f"{f(row, 'steady_rms_position_error_m'):.4f}",
            f"{f(row, 'final_position_error_m'):.4f}",
            f"{f(row, 'max_altitude_error_m'):.4f}",
            f"{f(row, 'mean_control_effort_rad_s'):.4f}",
            f"{f(row, 'max_tilt_command_rad'):.4f}",
            f"{f(row, 'rotor_saturation_rate'):.4f}",
        ]
        cells = table.add_row().cells
        for i, val in enumerate(vals):
            run = cells[i].paragraphs[0].add_run(val)
            set_run_font(run, size=7)


def write_docx(rows):
    checks = pass_fail(rows)
    overall = all(ok for _, ok in checks)
    temp_mpc = lookup(rows, "temperature", "mpc")
    temp_rl2 = lookup(rows, "temperature", "rl_v2")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RL-v2超越MPC控制策略报告")
    set_run_font(run, size=18, bold=True, color=(31, 78, 121))
    add_para(doc, f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}", size=9, color=(102, 102, 102), align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "1. 目标与方法")
    add_para(doc, "本报告验证 RL-v2 前瞻残差强化学习控制器是否在匀速圆周抗扰任务中优于线性 MPC。RL-v2 保留稳定底层控制器，并加入未来参考点、环境扰动和旋翼余量信息。")
    add_para(doc, "训练流程包含 MPC imitation 和 CEM 直接策略搜索；训练过程中会打开实时进度窗口显示候选策略 cost、RMS 与 MPC 基准。")

    add_heading(doc, "2. 指标汇总")
    add_metrics_table(doc, rows)
    add_para(doc, f"温度扰动下，MPC RMS 为 {f(temp_mpc, 'rms_position_error_m'):.4f} m，RL-v2 RMS 为 {f(temp_rl2, 'rms_position_error_m'):.4f} m。", bold=True, color=(31, 78, 121))

    add_heading(doc, "3. 验收结论")
    add_para(doc, "总体结果：" + ("RL-v2 已达到优于 MPC 的验收标准。" if overall else "RL-v2 尚未完全达到优于 MPC 的验收标准。"), bold=True, color=(31, 78, 121) if overall else (180, 0, 0))
    for name, ok in checks:
        add_para(doc, f"{'PASS' if ok else 'FAIL'}：{name}", color=(0, 100, 0) if ok else (180, 0, 0))

    add_heading(doc, "4. 图表")
    add_picture(doc, "rl_v2_benchmark_trajectory_3d.png", "图1 六控制器三维轨迹对比")
    add_picture(doc, "rl_v2_benchmark_position_error.png", "图2 六控制器位置误差时序")
    add_picture(doc, "rl_v2_benchmark_metric_bars.png", "图3 RMS、稳定段 RMS 和终端误差")
    add_picture(doc, "rl_v2_benchmark_effort_feasibility.png", "图4 控制努力、倾角和旋翼饱和率")

    out_path = REPORT_DIR / f"{STEM}.docx"
    doc.save(out_path)
    return out_path


def export_pdf(docx_path):
    ps1 = ROOT / "scripts" / "reporting" / "export_docx_pdf.ps1"
    pdf_path = REPORT_DIR / f"{STEM}.pdf"
    if not ps1.exists():
        return
    subprocess.run(
        [
            "powershell",
            "-ExecutionPolicy",
            "Bypass",
            "-File",
            str(ps1),
            "-DocxPath",
            str(docx_path),
            "-PdfPath",
            str(pdf_path),
        ],
        check=False,
    )


def main():
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    rows = read_metrics()
    write_markdown(rows)
    docx_path = write_docx(rows)
    export_pdf(docx_path)
    print(REPORT_DIR / f"{STEM}.md")
    print(docx_path)
    print(REPORT_DIR / f"{STEM}.pdf")


if __name__ == "__main__":
    main()
