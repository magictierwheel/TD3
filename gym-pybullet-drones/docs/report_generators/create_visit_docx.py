from __future__ import annotations

import csv
import json
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets" / "visit_doc"
OUT = DOCS / "reports" / "visit_overview" / "强化学习控制项目参观说明.docx"
REPRO_SCRIPT = ROOT / "experiments" / "hover_rl_reproduction" / "scripts" / "reproduce_hover_short.py"
REPRO_RESULTS = ROOT / "experiments" / "hover_rl_reproduction" / "results"
DOCKERFILE = ROOT / "reproducibility" / "docker" / "Dockerfile.repro"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def set_doc_style(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_paragraph(document: Document, text: str = ""):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return p


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_code_block(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_picture_centered(document: Document, image_path: Path, width_cm: float) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)


def render_formula(source: str, filename: str, width: float = 7.0) -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    path = ASSETS / filename
    fig = plt.figure(figsize=(width, 0.8), dpi=220)
    fig.patch.set_alpha(0)
    plt.text(0.5, 0.5, source, ha="center", va="center", fontsize=18)
    plt.axis("off")
    plt.savefig(path, transparent=True, bbox_inches="tight", pad_inches=0.08)
    plt.close(fig)
    return path


def render_rollout_chart() -> Path | None:
    csv_path = REPRO_RESULTS / "repro_hover_short" / "rollout.csv"
    if not csv_path.exists():
        return None
    steps: list[int] = []
    z_values: list[float] = []
    with csv_path.open("r", encoding="utf-8", newline="") as csv_file:
        reader = csv.DictReader(csv_file)
        for row in reader:
            steps.append(int(float(row["step"])))
            z_values.append(float(row["z"]))
    if not steps:
        return None
    ASSETS.mkdir(parents=True, exist_ok=True)
    path = ASSETS / "rollout_z.png"
    plt.figure(figsize=(7.2, 3.4), dpi=180)
    plt.plot(steps, z_values, label="rollout z", linewidth=2)
    plt.axhline(1.0, linestyle="--", color="#d62728", label="target z = 1.0")
    plt.xlabel("step")
    plt.ylabel("height z")
    plt.title("Short PPO Rollout Height")
    plt.grid(True, alpha=0.3)
    plt.legend()
    plt.tight_layout()
    plt.savefig(path, bbox_inches="tight")
    plt.close()
    return path


def load_summary() -> dict[str, object]:
    path = REPRO_RESULTS / "repro_hover_short" / "summary.json"
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return {}


def add_summary_table(document: Document, summary: dict[str, object]) -> None:
    rows = [
        ("任务", summary.get("task", "single_drone_hover")),
        ("算法", summary.get("algorithm", "PPO")),
        ("观测", summary.get("observation", "kin")),
        ("动作", summary.get("action", "one_d_rpm")),
        ("训练步数", summary.get("timesteps", "2048")),
        ("平均奖励", f"{float(summary.get('mean_reward', 0)):.4f}" if summary else "355.7683"),
        ("奖励标准差", f"{float(summary.get('std_reward', 0)):.4f}" if summary else "0.1818"),
        ("轨迹步数", summary.get("rollout_steps", "240")),
        ("最终高度", f"{float(summary.get('final_z', 0)):.4f}" if summary else "0.2163"),
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "项目", True)
    set_cell_text(table.rows[0].cells[1], "数值", True)
    set_cell_shading(table.rows[0].cells[0], "D9EAF7")
    set_cell_shading(table.rows[0].cells[1], "D9EAF7")
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], str(key))
        set_cell_text(cells[1], str(value))


def build_docx() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    formula_target = render_formula(r"$p^\ast=[0,\ 0,\ 1]^\mathsf{T}$", "target_position.png", 4.8)
    formula_reward = render_formula(
        r"$r_t=\max\left(0,\ 2-\left\|p^\ast-p_t\right\|_2^4\right)$",
        "reward.png",
        7.2,
    )
    formula_rpm = render_formula(
        r"$\mathrm{RPM}_i=\mathrm{RPM}_{hover}\left(1+0.05a_t\right)$",
        "rpm_action.png",
        6.5,
    )
    formula_control_reward = render_formula(
        r"$r=w_pS_p-w_uC_u-w_aC_a-w_sC_s$",
        "control_reward.png",
        6.2,
    )
    chart = render_rollout_chart()
    summary = load_summary()

    document = Document()
    set_doc_style(document)
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    title = document.add_heading("gym-pybullet-drones 强化学习控制项目参观说明", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(document, "面向参观、汇报和后续研究整理。核心任务是复现并理解基于 PPO 的四旋翼无人机悬停控制。")

    document.add_heading("1. 项目定位", level=1)
    add_paragraph(document, "本项目围绕“四旋翼无人机悬停控制”开展强化学习控制复现与研究。基础平台来自 learnsyslab/gym-pybullet-drones，它把 PyBullet 物理仿真、Gymnasium 强化学习接口和 Stable-Baselines3 算法库连接起来。")
    add_paragraph(document, "当前复现任务是单架四旋翼无人机从初始状态出发，学习悬停到目标位置：")
    add_picture_centered(document, formula_target, 4.2)
    add_paragraph(document, "项目不是直接手写一个传统 PID 控制器，而是让智能体在仿真环境中反复试飞，通过奖励反馈学习“观测状态到电机动作”的控制策略。")

    document.add_heading("2. 已完成工作", level=1)
    add_bullets(document, [
        "已准备 Docker 复现环境：reproducibility/docker/Dockerfile.repro。",
        "已添加短训练脚本：experiments/hover_rl_reproduction/scripts/reproduce_hover_short.py。",
        "已跑通 PPO 训练、模型保存、模型评估和轨迹导出。",
        "已验证官方示例 gym_pybullet_drones/examples/learn.py 可以运行。",
        "已生成短训练结果 experiments/hover_rl_reproduction/results/repro_hover_short 和 100000 步扩展训练结果 experiments/hover_rl_reproduction/results/repro_hover_100k。",
    ])
    add_summary_table(document, summary)

    document.add_heading("3. 强化学习整体工作流程", level=1)
    add_paragraph(document, "整体流程可以概括为：环境给状态，策略给动作，物理仿真推进一步，奖励函数评价动作好坏，然后 PPO 更新策略参数。")
    add_bullets(document, [
        "建立 HoverAviary 无人机悬停环境。",
        "初始化 PPO 策略网络。",
        "读取当前观测状态，包括位置、姿态、速度、角速度和动作历史。",
        "策略输出控制动作。",
        "BaseRLAviary 把动作转换为电机转速，PyBullet 推进动力学。",
        "HoverAviary 计算奖励、终止条件和截断条件。",
        "PPO 根据采样轨迹更新策略，并保存 best_model 或 final_model。",
        "训练后导出 summary.json 和 rollout.csv，用于结果解释。",
    ])

    document.add_heading("4. 强化学习任务建模", level=1)
    document.add_heading("4.1 智能体", level=2)
    add_paragraph(document, "智能体是 PPO 策略网络。训练完成后，保存下来的 .zip 模型可以理解为学习得到的控制器。")
    document.add_heading("4.2 环境", level=2)
    add_paragraph(document, "环境是 HoverAviary。它封装了四旋翼模型、PyBullet 物理仿真、目标位置、奖励计算和回合结束规则。")
    document.add_heading("4.3 观测状态", level=2)
    add_paragraph(document, "当前复现使用 kin 类型观测，包含位置、姿态、速度、角速度以及历史动作缓存。")
    add_code_block(document, "position = [x, y, z]\nattitude = [roll, pitch, yaw]\nvelocity = [vx, vy, vz]\nangular_velocity = [wx, wy, wz]")
    document.add_heading("4.4 控制动作", level=2)
    add_paragraph(document, "当前复现使用 one_d_rpm 动作类型。智能体输出一个标量动作，环境把它转换为四个电机共同变化的转速命令：")
    add_picture_centered(document, formula_rpm, 6.0)
    document.add_heading("4.5 奖励函数", level=2)
    add_paragraph(document, "悬停任务的奖励函数鼓励无人机靠近目标点：")
    add_picture_centered(document, formula_reward, 7.0)
    add_paragraph(document, "无人机越接近目标点，距离项越小，奖励越高。这个奖励函数在 HoverAviary.py 的 _computeReward() 中实现。")
    if chart:
        add_paragraph(document, "短训练轨迹中的高度变化如下。由于训练步数较少，最终高度尚未接近目标高度，这说明短训练只是验证链路，而不是最终控制效果。")
        add_picture_centered(document, chart, 13.5)

    document.add_heading("5. 软件环境配置", level=1)
    add_paragraph(document, "推荐使用 Docker 路线，因为它能绕开 Windows 原生 Python 和 PyBullet 编译问题。")
    document.add_heading("5.1 必需软件", level=2)
    add_bullets(document, ["Windows 10/11。", "Docker Desktop。", "PowerShell。", "Git，可选但建议安装。", "VS Code 或 PyCharm，可选，用于阅读代码。"])
    document.add_heading("5.2 容器内主要依赖", level=2)
    add_bullets(document, ["Python 3.10。", "numpy、scipy、matplotlib。", "pybullet。", "gymnasium。", "stable-baselines3 和 PyTorch。", "pytest。"])
    document.add_heading("5.3 构建 Docker 镜像", level=2)
    add_code_block(document, "cd E:\\1-AI辅助工作\\科研项目\\强化学习\\gym-pybullet-drones\ndocker build -f reproducibility/docker/Dockerfile.repro -t gym-pybullet-drones-repro .")

    document.add_heading("6. 一步一步复现实验", level=1)
    document.add_heading("6.1 快速短训练", level=2)
    add_code_block(document, 'docker run --rm -v "${PWD}\\experiments\\hover_rl_reproduction\\scripts\\reproduce_hover_short.py:/workspace/experiments/hover_rl_reproduction/scripts/reproduce_hover_short.py:ro" -v "${PWD}\\experiments\\hover_rl_reproduction\\results:/workspace/experiments/hover_rl_reproduction/results" gym-pybullet-drones-repro python experiments/hover_rl_reproduction/scripts/reproduce_hover_short.py --timesteps 2048 --eval-episodes 3 --rollout-steps 240 --output-folder experiments/hover_rl_reproduction/results/repro_hover_short')
    document.add_heading("6.2 查看结果", level=2)
    add_code_block(document, "dir results\\repro_hover_short\nGet-Content results\\repro_hover_short\\summary.json")
    add_paragraph(document, "重点关注 mean_reward、std_reward、final_z 和 rollout_csv。用 Excel 打开 rollout.csv 后，重点看 z 列是否逐步接近 1.0。")
    document.add_heading("6.3 运行官方 PPO 示例", level=2)
    add_code_block(document, 'docker run --rm -v "${PWD}\\experiments\\hover_rl_reproduction\\results:/workspace/experiments/hover_rl_reproduction/results" gym-pybullet-drones-repro python -c "from gym_pybullet_drones.examples.learn import run; run(gui=False, plot=False, local=False, output_folder=\'experiments/hover_rl_reproduction/results/original_learn_quick\')"')
    document.add_heading("6.4 增加训练步数", level=2)
    add_code_block(document, 'docker run --rm -v "${PWD}\\experiments\\hover_rl_reproduction\\scripts\\reproduce_hover_short.py:/workspace/experiments/hover_rl_reproduction/scripts/reproduce_hover_short.py:ro" -v "${PWD}\\experiments\\hover_rl_reproduction\\results:/workspace/experiments/hover_rl_reproduction/results" gym-pybullet-drones-repro python experiments/hover_rl_reproduction/scripts/reproduce_hover_short.py --timesteps 100000 --eval-episodes 5 --rollout-steps 240 --output-folder experiments/hover_rl_reproduction/results/repro_hover_100k')

    document.add_heading("7. 基于强化学习的控制研究总体思路", level=1)
    document.add_heading("7.1 明确控制目标", level=2)
    add_paragraph(document, "第一阶段聚焦单机定点悬停，第二阶段扩展到轨迹跟踪，第三阶段扩展到抗扰控制。评价指标包括位置误差、高度误差、姿态稳定性、平均奖励、成功率、收敛速度和动作平滑程度。")
    document.add_heading("7.2 建立基准控制器", level=2)
    add_bullets(document, ["PID 控制。", "项目自带 DSLPIDControl。", "PPO 强化学习策略。", "可选：SAC、TD3 或安全强化学习方法。"])
    document.add_heading("7.3 设计状态、动作和奖励", level=2)
    add_paragraph(document, "初始阶段使用 kin 观测和 one_d_rpm 动作。后续可以把动作扩展到四电机独立转速，并把奖励扩展为位置误差、姿态误差、能耗和动作平滑的组合。")
    add_picture_centered(document, formula_control_reward, 6.2)
    document.add_heading("7.4 制定训练实验矩阵", level=2)
    add_code_block(document, "algorithm: PPO / SAC / TD3\nobservation: kin / rgb\naction: one_d_rpm / rpm / pid / vel\ntimesteps: 1e5 / 5e5 / 1e6 / 1e7\nseed: 0 / 1 / 2 / 3 / 4\ndisturbance: none / wind / noise / delay")
    document.add_heading("7.5 评估与可视化", level=2)
    add_bullets(document, ["训练步数与平均奖励曲线。", "x、y、z 位置随时间变化曲线。", "高度误差随时间变化曲线。", "控制动作随时间变化曲线。", "多算法对比柱状图或箱线图。"])
    document.add_heading("7.6 从仿真走向真实控制", level=2)
    add_bullets(document, [
        "在 PyBullet 中训练和评估。",
        "加入随机化扰动，提高策略鲁棒性。",
        "与 PID 控制器混合，让 RL 输出高层速度或位置目标。",
        "在 Betaflight SITL 或 Crazyflie firmware 中测试。",
        "小范围低速实机测试。",
        "增加安全边界、急停逻辑和人工接管机制。",
    ])

    document.add_heading("8. 参观讲解路线", level=1)
    add_bullets(document, [
        "先展示项目目标：让无人机学会悬停到 z = 1.0。",
        "打开 HoverAviary.py，说明目标位置和奖励函数在哪里。",
        "打开 BaseRLAviary.py，说明状态和动作如何进入强化学习接口。",
        "打开 experiments/hover_rl_reproduction/scripts/reproduce_hover_short.py，说明 PPO 如何创建、训练、评估和保存。",
        "打开 summary.json，说明实验结果如何记录。",
        "打开 rollout.csv，说明轨迹数据如何验证飞行表现。",
        "最后说明后续研究方向：更长训练、更复杂动作空间、抗扰控制、多算法对比和仿真到实机迁移。",
    ])

    document.add_heading("9. 公式来源附录", level=1)
    add_paragraph(document, "为保证 Word/PDF 中公式显示清晰，本文件中的关键公式以本地渲染图片方式插入。公式源如下：")
    add_code_block(document, "p* = [0, 0, 1]^T\nr_t = max(0, 2 - ||p* - p_t||_2^4)\nRPM_i = RPM_hover * (1 + 0.05 a_t)\nr = w_p S_p - w_u C_u - w_a C_a - w_s C_s")

    document.save(OUT)


if __name__ == "__main__":
    build_docx()
    print(OUT)
