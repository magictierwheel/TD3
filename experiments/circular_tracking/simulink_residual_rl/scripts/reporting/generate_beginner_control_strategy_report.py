# -*- coding: utf-8 -*-
"""Generate a beginner-friendly report for the five quadrotor controllers."""

from __future__ import annotations

import csv
import json
import subprocess
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
REPORT_DIR = ROOT / "reports" / "strategy_comparison"
FIG_DIR = ROOT / "results" / "figures"
DATA_DIR = ROOT / "results" / "data"
FORMULA_DIR = ROOT / "reports" / "assets" / "beginner_control_strategies"
REPORT_STEM = "五种控制策略小白说明报告"
RENDER_SCRIPT = Path.home() / ".codex" / "skills" / "word-latex-academic-writing" / "scripts" / "render_formula.py"


CONTROLLER_ORDER = ["baseline", "pid_ff", "mpc", "adrc", "rl"]
MODEL_ORDER = ["standard", "temperature", "dust"]


FORMULAS = [
    {
        "id": "eq01",
        "title": "状态和目标信号",
        "formula": r"\mathbf{x}=\begin{bmatrix}\mathbf{r}\\ \mathbf{v}\\ \boldsymbol{\eta}\\ \boldsymbol{\omega}\\ \mathbf{e}_I\\ \boldsymbol{\Omega}\end{bmatrix},\qquad \mathbf{q}_{ref}=\begin{bmatrix}\mathbf{r}_{ref}\\ \mathbf{v}_{ref}\\ \mathbf{a}_{ref}\\ \psi_{ref}\end{bmatrix}",
        "explain": "x 是无人机当前状态，包含位置、速度、姿态、角速度、积分误差和四个旋翼实际转速。q_ref 是目标轨迹给控制器看的目标位置、目标速度、目标加速度和目标航向角。",
        "width": 5.9,
    },
    {
        "id": "eq02",
        "title": "匀速圆周参考轨迹",
        "formula": r"\mathbf{r}_{ref}(t)=\begin{bmatrix}R\cos(\omega_c t)\\ R\sin(\omega_c t)\\ h\end{bmatrix},\quad \mathbf{v}_{ref}(t)=\begin{bmatrix}-R\omega_c\sin(\omega_c t)\\ R\omega_c\cos(\omega_c t)\\0\end{bmatrix},\quad \mathbf{a}_{ref}(t)=\begin{bmatrix}-R\omega_c^2\cos(\omega_c t)\\-R\omega_c^2\sin(\omega_c t)\\0\end{bmatrix}",
        "explain": "这一组公式告诉无人机每一时刻应该在圆上的哪个点、应该朝哪个方向运动，以及圆周运动本身需要的向心加速度。",
        "width": 6.3,
    },
    {
        "id": "eq03",
        "title": "位置误差和速度误差",
        "formula": r"\mathbf{e}_r=\mathbf{r}_{ref}-\mathbf{r},\qquad \mathbf{e}_v=\mathbf{v}_{ref}-\mathbf{v}",
        "explain": "误差就是目标和实际之间的差。控制器看到的主要信息就是 e_r 和 e_v。",
        "width": 4.2,
    },
    {
        "id": "eq04",
        "title": "原 PID 外环",
        "formula": r"\mathbf{a}_{PID}=\mathbf{a}_{ref}+K_p\mathbf{e}_r+K_d\mathbf{e}_v+K_i\mathbf{e}_I",
        "explain": "原 PID 把偏离距离、偏离速度和长期累积小偏差合成一个期望加速度。这个加速度再被转换成无人机该倾斜多少、该给多少推力。",
        "width": 4.6,
    },
    {
        "id": "eq05",
        "title": "环境和旋翼效率",
        "formula": r"\rho=\frac{P}{R_{air}(T_0+\Delta T)},\qquad f_T=\max\left(0.25,\frac{\rho}{\rho_0}\eta_T\right),\qquad f_Q=\max\left(0.25,\frac{\rho}{\rho_0}\eta_Q\right)",
        "explain": "温度升高会让空气密度 rho 下降。粉尘会让 eta_T 和 eta_Q 变小。f_T 和 f_Q 表示旋翼产生推力和反扭矩的能力还剩多少。",
        "width": 5.8,
    },
    {
        "id": "eq06",
        "title": "风阻和热上升扰动",
        "formula": r"\mathbf{a}_{drag}=-\frac{1}{2m}\rho C_DA\,\|\mathbf{v}-\mathbf{w}\|(\mathbf{v}-\mathbf{w}),\qquad \mathbf{a}_{thermal}=\begin{bmatrix}0\\0\\a_{th}\end{bmatrix}",
        "explain": "风阻和热上升是外界直接推着无人机走的扰动。风阻跟相对风速有关，热上升主要影响竖直方向。",
        "width": 6.0,
    },
    {
        "id": "eq07",
        "title": "PID 扰动前馈补偿",
        "formula": r"\mathbf{a}_{cmd}=\mathbf{a}_{PID}-\mathbf{a}_{drag}-\mathbf{a}_{thermal},\qquad T_{cmd}\leftarrow \frac{T_{cmd}}{f_T},\qquad \boldsymbol{\tau}_{xy}\leftarrow \frac{\boldsymbol{\tau}_{xy}}{f_T},\quad \tau_z\leftarrow \frac{\tau_z}{f_Q}",
        "explain": "如果模型已经知道风阻、热上升和推力变弱，就可以提前补偿。箭头表示在原命令基础上做修正。",
        "width": 6.2,
    },
    {
        "id": "eq08",
        "title": "MPC 的预测模型和优化目标",
        "formula": r"\begin{aligned}\mathbf{s}_{k+1}&=\begin{bmatrix}1&T_s\\0&1\end{bmatrix}\mathbf{s}_k+\begin{bmatrix}\frac{1}{2}T_s^2\\T_s\end{bmatrix}u_k,\\ \min_{\{u_k\}}\ J&=\sum_{k=1}^{N}\left(q_r\|r_k-r_{ref,k}\|^2+q_v\|v_k-v_{ref,k}\|^2+r_u\|u_k\|^2\right)\end{aligned}",
        "explain": "MPC 先用简化模型预测未来 N 步，再选择一串加速度命令，使未来位置和速度尽量贴近目标，同时不要让控制动作太大。",
        "width": 6.1,
    },
    {
        "id": "eq09",
        "title": "ADRC 的扩张状态观测器",
        "formula": r"\begin{aligned}\dot{\mathbf{z}}_1&=\mathbf{z}_2-\beta_1(\mathbf{z}_1-\mathbf{r}),\\ \dot{\mathbf{z}}_2&=\mathbf{z}_3+\mathbf{u}-\beta_2(\mathbf{z}_1-\mathbf{r}),\\ \dot{\mathbf{z}}_3&=-\beta_3(\mathbf{z}_1-\mathbf{r})\end{aligned}",
        "explain": "z1 估计位置，z2 估计速度，z3 估计总扰动。ADRC 不要求先完全写出扰动公式，而是从运动偏差中估计扰动。",
        "width": 5.5,
    },
    {
        "id": "eq10",
        "title": "ADRC 外环控制律",
        "formula": r"\mathbf{u}_{ADRC}=\mathbf{a}_{ref}+K_p(\mathbf{r}_{ref}-\mathbf{r})+K_d(\mathbf{v}_{ref}-\mathbf{v})-\Gamma\mathbf{z}_3",
        "explain": "ADRC 仍然会按位置和速度误差纠正，但会额外减去估计到的总扰动。Gamma 表示补偿比例。",
        "width": 5.3,
    },
    {
        "id": "eq11",
        "title": "强化学习残差策略",
        "formula": r"\mathbf{a}_{cmd}=\mathbf{a}_{PID}+\Delta\mathbf{a}_{RL},\qquad [\Delta\mathbf{a}_{RL},s_T,s_{\tau}]=\pi_{\theta}(\mathbf{x},\mathbf{q}_{ref},\mathbf{e}_{env})",
        "explain": "RL 策略不是推倒重来，而是在原 PID 上加一个训练出来的补偿量。pi_theta 表示训练得到的策略。",
        "width": 5.8,
    },
    {
        "id": "eq12",
        "title": "加速度到目标姿态",
        "formula": r"\phi_{des}=\frac{a_x\sin\psi_{ref}-a_y\cos\psi_{ref}}{g},\qquad \theta_{des}=\frac{a_x\cos\psi_{ref}+a_y\sin\psi_{ref}}{g}",
        "explain": "四旋翼想往水平某个方向加速，就要朝那个方向倾斜。phi_des 和 theta_des 就是期望横滚角和期望俯仰角。",
        "width": 5.8,
    },
    {
        "id": "eq13",
        "title": "总推力命令",
        "formula": r"T_{cmd}=\frac{m(g+a_z)}{\cos\phi\cos\theta}",
        "explain": "无人机要抵消重力，还要按需要上升或下降。机身倾斜后，竖直方向能用到的推力会变少，所以分母里有 cos 项。",
        "width": 3.7,
    },
    {
        "id": "eq14",
        "title": "四旋翼电机分配",
        "formula": r"\begin{bmatrix}T\\\tau_x\\\tau_y\\\tau_z\end{bmatrix}=\begin{bmatrix}k_f&k_f&k_f&k_f\\0&lk_f&0&-lk_f\\-lk_f&0&lk_f&0\\-k_q&k_q&-k_q&k_q\end{bmatrix}\begin{bmatrix}\Omega_1^2\\\Omega_2^2\\\Omega_3^2\\\Omega_4^2\end{bmatrix}",
        "explain": "控制器最终需要的是四个电机转速。这个矩阵把总推力和三个方向的转动力矩分配给四个旋翼。",
        "width": 5.7,
    },
    {
        "id": "eq15",
        "title": "动力学闭环",
        "formula": r"\dot{\mathbf{r}}=\mathbf{v},\qquad \dot{\mathbf{v}}=\frac{R_{BW}\begin{bmatrix}0\\0\\T\end{bmatrix}}{m}-\begin{bmatrix}0\\0\\g\end{bmatrix}+\mathbf{a}_{drag}+\mathbf{a}_{thermal}",
        "explain": "这是仿真中无人机真正运动的更新规则。控制器给推力和姿态，环境给扰动，动力学把它们合成新的位置和速度。",
        "width": 6.0,
    },
]


def set_run_font(run, size=11, bold=False, color=None, name="Microsoft YaHei"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)


def add_para(doc, text="", size=11, bold=False, color=None, align=None):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_heading(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=14, bold=True, color=(31, 78, 121))
    return paragraph


def add_picture(doc, image_name, caption, width=6.2):
    image_path = FIG_DIR / image_name
    if not image_path.exists():
        add_para(doc, f"[缺失图片: {image_name}]", color=(180, 0, 0))
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run(caption)
    set_run_font(caption_run, size=9, color=(102, 102, 102))


def render_formulas():
    FORMULA_DIR.mkdir(parents=True, exist_ok=True)
    sources = ["# Formula Sources\n"]
    for index, spec in enumerate(FORMULAS, start=1):
        png_path = FORMULA_DIR / f"{spec['id']}.png"
        cmd = [
            sys.executable,
            str(RENDER_SCRIPT),
            "--formula",
            spec["formula"],
            "--display",
            "--out-dir",
            str(FORMULA_DIR),
            "--name",
            png_path.name,
            "--font-size",
            "12",
            "--dpi",
            "420",
        ]
        result = subprocess.run(
            cmd,
            cwd=str(ROOT),
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            encoding="utf-8",
            errors="replace",
            check=False,
        )
        if result.returncode != 0:
            raise RuntimeError(f"Formula rendering failed for {spec['id']}:\n{result.stdout}")
        try:
            meta = json.loads(result.stdout)
            source = meta.get("source", spec["formula"])
        except json.JSONDecodeError:
            source = spec["formula"]
        sources.extend(
            [
                f"## ({index}) {spec['title']}",
                "",
                "```latex",
                source,
                "```",
                "",
            ]
        )
    (FORMULA_DIR / "formula_sources.md").write_text("\n".join(sources), encoding="utf-8")


def add_formula_image(doc, spec, number):
    image_path = FORMULA_DIR / f"{spec['id']}.png"
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(spec.get("width", 5.5)))
    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = label.add_run(f"式（{number}）{spec['title']}")
    set_run_font(label_run, size=9, color=(102, 102, 102))
    add_para(doc, spec["explain"], size=10)


def read_metrics():
    path = DATA_DIR / "quadrotor_strategy_circle_comparison_metrics.csv"
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def f(row, key):
    return float(row[key])


def lookup(rows, model_type, controller_type):
    for row in rows:
        if row["model_type"] == model_type and row["controller_type"] == controller_type:
            return row
    raise KeyError((model_type, controller_type))


def reduction(base, candidate, key="rms_position_error_m"):
    b = f(base, key)
    c = f(candidate, key)
    if abs(b) < 1.0e-12:
        return 0.0
    return 100.0 * (b - c) / b


def controller_table_markdown(rows):
    lines = [
        "| 控制策略 | 标准 RMS/m | 温度扰动 RMS/m | 粉尘扰动 RMS/m | 直观理解 |",
        "|---|---:|---:|---:|---|",
    ]
    notes = {
        "baseline": "发现偏了再拉回来",
        "pid_ff": "提前知道风和效率损失，先补一把",
        "mpc": "每一小步都看未来轨迹再决定",
        "adrc": "边飞边估计看不见的推力和扰动",
        "rl": "把仿真训练出的补偿经验加到原控制器上",
    }
    for controller in CONTROLLER_ORDER:
        std = lookup(rows, "standard", controller)
        temp = lookup(rows, "temperature", controller)
        dust = lookup(rows, "dust", controller)
        lines.append(
            f"| {std['controller_label']} | "
            f"{f(std, 'rms_position_error_m'):.4f} | "
            f"{f(temp, 'rms_position_error_m'):.4f} | "
            f"{f(dust, 'rms_position_error_m'):.4f} | "
            f"{notes[controller]} |"
        )
    return "\n".join(lines)


def beginner_content(rows):
    temp_base = lookup(rows, "temperature", "baseline")
    dust_base = lookup(rows, "dust", "baseline")
    std_base = lookup(rows, "standard", "baseline")
    temp_pidff = lookup(rows, "temperature", "pid_ff")
    temp_mpc = lookup(rows, "temperature", "mpc")
    temp_adrc = lookup(rows, "temperature", "adrc")
    temp_rl = lookup(rows, "temperature", "rl")
    dust_pidff = lookup(rows, "dust", "pid_ff")
    dust_mpc = lookup(rows, "dust", "mpc")
    dust_adrc = lookup(rows, "dust", "adrc")
    dust_rl = lookup(rows, "dust", "rl")
    std_adrc = lookup(rows, "standard", "adrc")

    return {
        "title": "五种控制策略小白说明报告",
        "generated": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "summary": [
            "这份报告不从公式开始，而是从一个问题开始：无人机想沿着一个圆飞，风、热上升和粉尘让它偏离路线时，控制器怎样把它拉回去。",
            "本项目里的五种控制策略都做同一件事：看无人机现在在哪里、应该在哪里，然后把纠正动作变成四个旋翼的转速。差别在于它们判断“该怎么纠正”的方式不同。",
        ],
        "key_numbers": [
            f"温度扰动下，原 PID 的全程 RMS 误差为 {f(temp_base, 'rms_position_error_m'):.4f} m；PID 补偿、MPC、ADRC、RL 分别为 {f(temp_pidff, 'rms_position_error_m'):.4f} m、{f(temp_mpc, 'rms_position_error_m'):.4f} m、{f(temp_adrc, 'rms_position_error_m'):.4f} m、{f(temp_rl, 'rms_position_error_m'):.4f} m。",
            f"粉尘扰动下，原 PID 的全程 RMS 误差为 {f(dust_base, 'rms_position_error_m'):.4f} m；ADRC 降到 {f(dust_adrc, 'rms_position_error_m'):.4f} m，约降低 {reduction(dust_base, dust_adrc):.1f}%。",
            f"标准环境下，ADRC 的全程 RMS 为 {f(std_adrc, 'rms_position_error_m'):.4f} m，略低于原 PID 的 {f(std_base, 'rms_position_error_m'):.4f} m。",
        ],
        "table": controller_table_markdown(rows),
        "reductions": {
            "pidff_temp": reduction(temp_base, temp_pidff),
            "mpc_temp": reduction(temp_base, temp_mpc),
            "adrc_temp": reduction(temp_base, temp_adrc),
            "rl_temp": reduction(temp_base, temp_rl),
            "pidff_dust": reduction(dust_base, dust_pidff),
            "mpc_dust": reduction(dust_base, dust_mpc),
            "adrc_dust": reduction(dust_base, dust_adrc),
            "rl_dust": reduction(dust_base, dust_rl),
        },
    }


def write_markdown(rows):
    content = beginner_content(rows)
    report_path = REPORT_DIR / f"{REPORT_STEM}.md"
    lines = [
        f"# {content['title']}",
        "",
        f"生成时间：{content['generated']}",
        "",
        "## 1. 先用一句话理解无人机控制",
        "",
        "无人机控制就像人端着一杯水走圆形路线。目标不是只往前走，而是每一瞬间都要知道自己离路线偏了多少、偏离趋势有多快、要不要提前用力修正。",
        "",
        "在这个项目里，圆形路线由“参考轨迹”给出。仿真模型会不断计算无人机的真实位置、速度、姿态和旋翼转速。控制器看到“真实状态”和“目标状态”的差距后，会输出四个电机应该转多快。",
        "",
        "可以把整个过程理解成下面这条链：",
        "",
        "```text",
        "目标圆形路线 -> 比较当前位置 -> 控制器决定怎么纠正 -> 换成总推力和姿态力矩 -> 分配到四个旋翼 -> 无人机运动改变",
        "```",
        "",
        "## 2. 本报告里的几个简单词",
        "",
        "- 误差：无人机真实位置和目标位置之间的距离。",
        "- RMS 误差：把整段飞行的偏离程度合成一个数。越小，说明整体越贴近目标圆。",
        "- 稳定段 RMS：仿真 8 秒之后的 RMS。它更像“进入稳定飞行后还偏多少”。",
        "- 控制努力：电机转速调整有多忙。数值越高，说明控制器更积极，也可能更耗能或更抖。",
        "- 扰动：外界或执行器造成的麻烦。本项目主要有温度扰动、热上升、风和粉尘导致的推力效率下降。",
        "",
        "## 3. 五种控制策略是怎么起作用的",
        "",
        "### 3.1 原 PID：发现偏了，再按偏差拉回来",
        "",
        "原 PID 是最朴素的办法。它像一个人看到自己走偏了，就朝反方向调整。偏得越远，拉回来的动作越大；偏离速度越快，刹车或纠正也越明显；如果长期有一点点偏差，还会慢慢补上。",
        "",
        "在本项目中，原 PID 先把位置误差变成想要的加速度，再换成目标俯仰角、横滚角、总推力和力矩，最后变成四个旋翼转速。它在标准环境下能稳定圆周飞行，但它不知道风、热上升或粉尘已经改变了飞行条件，所以扰动一来，误差会变大。",
        "",
        "### 3.2 PID + 扰动前馈：知道有风，就提前补偿",
        "",
        "PID 补偿版仍然保留原来的 PID，但多了一层“提前量”。好比你骑车时看到前面有侧风，不等车歪了才修正，而是先把车把稍微往反方向压一点。",
        "",
        "本项目中的环境模块会给出空气密度、热上升、风速、推力效率和反扭矩效率。PID 补偿控制器用这些信息抵消风阻和热上升，并在粉尘导致推力变弱时把电机推力命令放大一些。它的优点是直观、好解释；缺点是补偿效果依赖扰动模型是否写得准。",
        "",
        "### 3.3 线性 MPC：每一步都看一小段未来",
        "",
        "MPC 可以理解成“会预判的司机”。普通 PID 主要看现在偏了多少，MPC 会看接下来一小段圆形路线，然后选择一个看起来最合适的加速度命令。",
        "",
        "本项目里的 MPC 使用简化的直线运动近似来预测未来位置和速度，并限制加速度、倾角和推力。它仍然沿用后面的姿态控制和旋翼分配链路。MPC 在温度扰动下误差最低，但它的控制努力也更高，说明它为了追踪得更紧，电机调整更积极。",
        "",
        "### 3.4 ADRC：边飞边估计“看不见的推手”",
        "",
        "ADRC 的核心想法是：不一定非要提前知道所有扰动公式，只要能从无人机的运动变化中估计“是不是有什么东西在推它”。这个估计器叫 ESO，可以把它想成一个观察员。观察员不直接看风，也不直接看热气流，而是看无人机有没有出现不该出现的加速或偏移。",
        "",
        "本项目中的 ADRC 用 ESO 估计总扰动，再把估计到的扰动从加速度命令里抵消掉。同时，粉尘会让旋翼推力效率下降，这属于执行器变弱，因此 ADRC 也使用推力/力矩效率调度来修正电机命令。结果上，ADRC 在标准和粉尘工况下表现很好；在温度扰动下也优于原 PID，但不如显式补偿、MPC 和 RL。",
        "",
        "### 3.5 强化学习策略：把训练出来的补偿经验加到原控制器上",
        "",
        "强化学习可以理解成“反复练习后形成的经验”。它不是完全替代底层控制器，而是在原控制器保持稳定的基础上，额外给出一个残差补偿动作。",
        "",
        "本项目里的 RL 策略看当前状态、目标轨迹和环境量，输出对加速度、推力和力矩的补偿。它不需要人工把每一种扰动都写成完整公式，但它的表现依赖训练范围。当前结果表明，RL 在温度和粉尘扰动下接近 PID 补偿效果，并保持标准工况不被破坏。",
        "",
        "## 4. 完整公式版全过程",
        "",
        "下面的公式不是要求零基础读者一次看懂所有符号，而是把“控制器到底怎么算”的完整链条保留下来。读法可以很简单：先看公式标题，再看公式下面那句中文解释。",
        "",
    ]
    for index, spec in enumerate(FORMULAS, start=1):
        lines.extend(
            [
                f"### 式（{index}）{spec['title']}",
                "",
                "$$",
                spec["formula"],
                "$$",
                "",
                spec["explain"],
                "",
            ]
        )
    lines.extend(
        [
        "## 5. 项目结果用小白语言怎么读",
        "",
        content["table"],
        "",
        "表里的 RMS 误差越小越好。它可以理解为“整段飞行平均偏离目标圆有多远”。",
        "",
        "## 6. 哪种策略适合什么场景",
        "",
        "- 原 PID：适合做基线。它简单、稳定，但不知道环境变化。",
        "- PID 扰动补偿：适合扰动模型比较清楚的情况，比如你能估计风、热上升和推力效率损失。",
        "- MPC：适合需要考虑未来轨迹和约束的情况。它追踪能力强，但电机调整更积极。",
        "- ADRC：适合扰动不容易精确建模、但希望用传统控制方法在线估计扰动的情况。",
        "- 强化学习策略：适合扰动复杂、难以完全写公式，但可以通过仿真训练积累补偿经验的情况。",
        "",
        "## 7. 这次仿真的主要结论",
        "",
        ]
    )
    for item in content["key_numbers"]:
        lines.append(f"- {item}")
    lines.extend(
        [
            f"- 温度扰动下，MPC 的误差相对原 PID 约降低 {content['reductions']['mpc_temp']:.1f}%，PID 补偿约降低 {content['reductions']['pidff_temp']:.1f}%，RL 约降低 {content['reductions']['rl_temp']:.1f}%，ADRC 约降低 {content['reductions']['adrc_temp']:.1f}%。",
            f"- 粉尘扰动下，ADRC 约降低 {content['reductions']['adrc_dust']:.1f}%，PID 补偿约降低 {content['reductions']['pidff_dust']:.1f}%，RL 约降低 {content['reductions']['rl_dust']:.1f}%。",
            "",
            "这说明五种方法不是简单的“谁永远最好”。原 PID 是参照物；PID 补偿和 MPC 在已知扰动下很有效；ADRC 提供了一条传统抗扰路线；RL 适合把复杂环境中的补偿经验训练出来。",
            "",
            "## 8. 配套图表",
            "",
            "- `results/figures/strategy_circle_trajectory_3d.png`：五种策略在三类环境下的圆周轨迹。",
            "- `results/figures/strategy_circle_position_error.png`：位置误差随时间变化。",
            "- `results/figures/strategy_circle_metric_bars.png`：RMS、稳定段 RMS、终端误差对比。",
            "- `results/figures/strategy_circle_effort_altitude.png`：控制努力和最大高度误差对比。",
            "",
            "![五种策略轨迹对比](../../results/figures/strategy_circle_trajectory_3d.png)",
            "",
            "![五种策略误差指标](../../results/figures/strategy_circle_metric_bars.png)",
        ]
    )
    report_path.write_text("\n".join(lines), encoding="utf-8-sig")
    return report_path


def add_docx_table(doc, rows):
    headers = ["控制策略", "标准 RMS/m", "温度 RMS/m", "粉尘 RMS/m", "直观理解"]
    notes = {
        "baseline": "发现偏了再拉回来",
        "pid_ff": "提前知道风和效率损失，先补一把",
        "mpc": "看未来一小段轨迹再决定",
        "adrc": "估计看不见的扰动并抵消",
        "rl": "把训练出的补偿经验加上去",
    }
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        run = table.rows[0].cells[i].paragraphs[0].add_run(header)
        set_run_font(run, size=8, bold=True)
    for controller in CONTROLLER_ORDER:
        std = lookup(rows, "standard", controller)
        temp = lookup(rows, "temperature", controller)
        dust = lookup(rows, "dust", controller)
        values = [
            std["controller_label"],
            f"{f(std, 'rms_position_error_m'):.4f}",
            f"{f(temp, 'rms_position_error_m'):.4f}",
            f"{f(dust, 'rms_position_error_m'):.4f}",
            notes[controller],
        ]
        cells = table.add_row().cells
        for i, value in enumerate(values):
            run = cells[i].paragraphs[0].add_run(value)
            set_run_font(run, size=8)


def write_docx(rows):
    content = beginner_content(rows)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(content["title"])
    set_run_font(run, size=18, bold=True, color=(31, 78, 121))
    add_para(doc, f"生成时间：{content['generated']}", size=9, color=(102, 102, 102), align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "1. 一句话理解无人机控制")
    add_para(doc, "无人机控制就像人端着一杯水走圆形路线。它每一瞬间都要知道自己离路线偏了多少、偏离趋势有多快、要不要提前用力修正。")
    add_para(doc, "在本项目中，控制链路是：目标圆形路线 -> 比较当前位置 -> 控制器决定纠正动作 -> 变成总推力和姿态力矩 -> 分配到四个旋翼。")

    add_heading(doc, "2. 五种控制策略")
    add_para(doc, "原 PID：发现偏了，再按偏差拉回来。它简单稳定，但不知道风、热上升或粉尘已经改变了飞行条件。")
    add_para(doc, "PID 扰动补偿：在 PID 基础上提前补偿风阻、热上升、密度变化和粉尘导致的效率损失。它好解释，但依赖扰动模型。")
    add_para(doc, "线性 MPC：每一步都看未来一小段圆形路线，再选一个合适的加速度命令。它追踪紧，但控制努力更高。")
    add_para(doc, "ADRC：用 ESO 估计看不见的总扰动，并对粉尘造成的执行器效率下降做推力/力矩调度。它在标准和粉尘工况下表现好，温度扰动下优于原 PID。")
    add_para(doc, "强化学习策略：把训练出来的补偿经验加到原控制器上。它不需要人工写全所有扰动公式，但依赖训练范围。")

    add_heading(doc, "3. 完整公式版全过程")
    add_para(doc, "下面的公式保留完整计算链条。阅读时先看公式标题，再看公式下方的中文解释；不需要先掌握所有符号。")
    for index, spec in enumerate(FORMULAS, start=1):
        add_formula_image(doc, spec, index)

    add_heading(doc, "4. 结果怎么读")
    add_para(doc, "RMS 误差越小，说明整段飞行越贴近目标圆。下面的表只保留最容易理解的全程 RMS。")
    add_docx_table(doc, rows)
    for item in content["key_numbers"]:
        add_para(doc, item, bold=True, color=(31, 78, 121))

    add_heading(doc, "5. 结论")
    add_para(doc, "原 PID 适合作为基线。PID 补偿和 MPC 在已知扰动下效果更好。ADRC 提供了一条传统抗扰路线，特别适合不想完全依赖扰动公式的场景。强化学习策略适合在复杂扰动中学习补偿经验。")

    add_heading(doc, "6. 图表")
    add_picture(doc, "strategy_circle_trajectory_3d.png", "图1 五种控制策略圆周轨迹对比")
    add_picture(doc, "strategy_circle_metric_bars.png", "图2 五种控制策略误差指标对比")

    out_path = REPORT_DIR / f"{REPORT_STEM}.docx"
    doc.save(out_path)
    return out_path


def main():
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    rows = read_metrics()
    render_formulas()
    md_path = write_markdown(rows)
    docx_path = write_docx(rows)
    print(md_path)
    print(docx_path)


if __name__ == "__main__":
    main()
